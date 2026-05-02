"""
Сервис формирования отчетов.

Модуль отвечает за:
- формирование DOCX-отчета по KPI сотрудников;
- формирование XLSX-отчета с аналитикой KPI;
- сохранение файлов в папку exports;
- запись информации о сформированных отчетах в базу данных.
"""

from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import List

from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment
from sqlalchemy.orm import joinedload

from app.database.connection import BASE_DIR
from app.database.models import Employee, KpiRecord, Report
from app.services.analytics_service import AnalyticsService
from app.services.auth_service import AuthenticatedUser


EXPORTS_DIR = BASE_DIR / "exports"
EXPORTS_DIR.mkdir(exist_ok=True)


@dataclass
class GeneratedReportResult:
    """
    Результат формирования отчета.
    """

    report_name: str
    report_type: str
    file_path: str


@dataclass
class ReportListItem:
    """
    Краткая информация о сформированном отчете.
    """

    id: int
    report_name: str
    report_type: str
    file_path: str
    created_at: datetime


class ReportService:
    """
    Сервис формирования отчетов.
    """

    def __init__(self, session):
        """
        Создает сервис отчетов.

        Параметры:
            session: активная сессия SQLAlchemy.
        """
        self.session = session

    def generate_kpi_docx_report(
        self,
        current_user: AuthenticatedUser,
    ) -> GeneratedReportResult:
        """
        Формирует DOCX-отчет по KPI сотрудников.

        В отчет включаются:
        - сведения о пользователе, сформировавшем отчет;
        - список доступных сотрудников;
        - KPI-записи сотрудников.
        """
        employees = self._get_allowed_employees(current_user)

        document = Document()

        document.add_heading("Отчет по KPI сотрудников", level=1)

        document.add_paragraph(
            "Система: KPI Monitor ML — система мониторинга и анализа KPI "
            "сотрудников организации с использованием методов машинного обучения."
        )
        document.add_paragraph(f"Сформировал: {current_user.full_name}")
        document.add_paragraph(f"Роль пользователя: {current_user.role_name}")
        document.add_paragraph(
            f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        )

        document.add_heading("Сотрудники", level=2)

        for employee in employees:
            document.add_heading(employee.full_name, level=3)
            document.add_paragraph(f"Табельный номер: {employee.personnel_number}")
            document.add_paragraph(
                f"Отдел: {employee.department.name if employee.department else '-'}"
            )
            document.add_paragraph(
                f"Должность: {employee.position.name if employee.position else '-'}"
            )

            records = (
                self.session.query(KpiRecord)
                .options(joinedload(KpiRecord.indicator))
                .filter(KpiRecord.employee_id == employee.id)
                .order_by(KpiRecord.period_start.desc())
                .all()
            )

            if not records:
                document.add_paragraph("KPI-записи отсутствуют.")
                continue

            table = document.add_table(rows=1, cols=6)
            table.style = "Table Grid"

            header_cells = table.rows[0].cells
            header_cells[0].text = "Показатель"
            header_cells[1].text = "Период с"
            header_cells[2].text = "Период по"
            header_cells[3].text = "Факт"
            header_cells[4].text = "План"
            header_cells[5].text = "Оценка"

            for record in records:
                row_cells = table.add_row().cells
                row_cells[0].text = record.indicator.name if record.indicator else "-"
                row_cells[1].text = record.period_start.strftime("%d.%m.%Y")
                row_cells[2].text = record.period_end.strftime("%d.%m.%Y")
                row_cells[3].text = f"{record.actual_value:.2f}"
                row_cells[4].text = f"{record.target_value:.2f}"
                row_cells[5].text = f"{record.score:.2f}"

        file_name = self._build_file_name("kpi_report", "docx")
        file_path = EXPORTS_DIR / file_name

        document.save(file_path)

        self._save_report_record(
            report_name="DOCX-отчет по KPI сотрудников",
            report_type="DOCX",
            file_path=file_path,
            current_user=current_user,
        )

        return GeneratedReportResult(
            report_name="DOCX-отчет по KPI сотрудников",
            report_type="DOCX",
            file_path=str(file_path),
        )

    def generate_analytics_xlsx_report(
        self,
        current_user: AuthenticatedUser,
    ) -> GeneratedReportResult:
        """
        Формирует XLSX-отчет с аналитикой KPI.

        В файл включаются:
        - сводные показатели;
        - аналитика по отделам;
        - аналитика по сотрудникам.
        """
        analytics_service = AnalyticsService(self.session)

        dashboard_stats = analytics_service.get_dashboard_stats(current_user)
        department_summary = analytics_service.get_department_summary(current_user)
        employee_summary = analytics_service.get_employee_summary(current_user)

        workbook = Workbook()

        summary_sheet = workbook.active
        summary_sheet.title = "Сводка"

        summary_sheet["A1"] = "Отчет аналитики KPI"
        summary_sheet["A1"].font = Font(bold=True, size=14)

        summary_rows = [
            ("Сформировал", current_user.full_name),
            ("Роль", current_user.role_name),
            ("Дата формирования", datetime.now().strftime("%d.%m.%Y %H:%M")),
            ("Количество сотрудников", dashboard_stats.employees_count),
            ("Количество KPI-записей", dashboard_stats.kpi_records_count),
            ("Средняя оценка", dashboard_stats.average_score),
            ("Минимальная оценка", dashboard_stats.min_score),
            ("Максимальная оценка", dashboard_stats.max_score),
        ]

        for row_index, row_data in enumerate(summary_rows, start=3):
            summary_sheet.cell(row=row_index, column=1, value=row_data[0])
            summary_sheet.cell(row=row_index, column=2, value=row_data[1])

        departments_sheet = workbook.create_sheet("По отделам")
        departments_headers = [
            "Отдел",
            "Количество сотрудников",
            "Количество KPI-записей",
            "Средняя оценка",
        ]
        self._write_headers(departments_sheet, departments_headers)

        for row_index, item in enumerate(department_summary, start=2):
            departments_sheet.cell(row=row_index, column=1, value=item.department_name)
            departments_sheet.cell(row=row_index, column=2, value=item.employees_count)
            departments_sheet.cell(row=row_index, column=3, value=item.records_count)
            departments_sheet.cell(row=row_index, column=4, value=item.average_score)

        employees_sheet = workbook.create_sheet("По сотрудникам")
        employees_headers = [
            "ID",
            "Сотрудник",
            "Отдел",
            "Количество KPI-записей",
            "Средняя оценка",
        ]
        self._write_headers(employees_sheet, employees_headers)

        for row_index, item in enumerate(employee_summary, start=2):
            employees_sheet.cell(row=row_index, column=1, value=item.employee_id)
            employees_sheet.cell(row=row_index, column=2, value=item.employee_name)
            employees_sheet.cell(row=row_index, column=3, value=item.department_name)
            employees_sheet.cell(row=row_index, column=4, value=item.records_count)
            employees_sheet.cell(row=row_index, column=5, value=item.average_score)

        for sheet in workbook.worksheets:
            self._auto_size_columns(sheet)

        file_name = self._build_file_name("kpi_analytics", "xlsx")
        file_path = EXPORTS_DIR / file_name

        workbook.save(file_path)

        self._save_report_record(
            report_name="XLSX-отчет аналитики KPI",
            report_type="XLSX",
            file_path=file_path,
            current_user=current_user,
        )

        return GeneratedReportResult(
            report_name="XLSX-отчет аналитики KPI",
            report_type="XLSX",
            file_path=str(file_path),
        )

    def get_reports(self) -> List[ReportListItem]:
        """
        Возвращает список сформированных отчетов.
        """
        reports = (
            self.session.query(Report)
            .order_by(Report.created_at.desc())
            .all()
        )

        return [
            ReportListItem(
                id=report.id,
                report_name=report.report_name,
                report_type=report.report_type,
                file_path=report.file_path,
                created_at=report.created_at,
            )
            for report in reports
        ]

    def _get_allowed_employees(
        self,
        current_user: AuthenticatedUser,
    ) -> List[Employee]:
        """
        Возвращает сотрудников, доступных пользователю по роли.
        """
        query = (
            self.session.query(Employee)
            .options(
                joinedload(Employee.department),
                joinedload(Employee.position),
            )
            .order_by(Employee.full_name)
        )

        if current_user.role_name == "Администратор":
            return query.all()

        if current_user.role_name == "Сотрудник":
            if current_user.employee_id is None:
                return []

            return query.filter(Employee.id == current_user.employee_id).all()

        if current_user.role_name == "Руководитель":
            if current_user.employee_id is None:
                return []

            manager_employee = (
                self.session.query(Employee)
                .filter(Employee.id == current_user.employee_id)
                .first()
            )

            if manager_employee is None:
                return []

            return query.filter(
                Employee.department_id == manager_employee.department_id
            ).all()

        return []

    def _save_report_record(
        self,
        report_name: str,
        report_type: str,
        file_path: Path,
        current_user: AuthenticatedUser,
    ) -> None:
        """
        Сохраняет информацию о сформированном отчете в базу данных.
        """
        report = Report(
            report_name=report_name,
            report_type=report_type,
            file_path=str(file_path),
            created_by_user_id=current_user.id,
        )

        self.session.add(report)
        self.session.commit()

    def _build_file_name(self, prefix: str, extension: str) -> str:
        """
        Формирует уникальное имя файла отчета.
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        return f"{prefix}_{timestamp}.{extension}"

    def _write_headers(self, sheet, headers: List[str]) -> None:
        """
        Записывает заголовки таблицы в Excel-лист.
        """
        for column_index, header in enumerate(headers, start=1):
            cell = sheet.cell(row=1, column=column_index, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal="center")

    def _auto_size_columns(self, sheet) -> None:
        """
        Автоматически подбирает ширину колонок Excel-листа.
        """
        for column_cells in sheet.columns:
            max_length = 0
            column_letter = column_cells[0].column_letter

            for cell in column_cells:
                cell_value = str(cell.value) if cell.value is not None else ""
                max_length = max(max_length, len(cell_value))

            sheet.column_dimensions[column_letter].width = max_length + 3